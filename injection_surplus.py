# -*- coding: utf-8 -*-

from datetime import datetime
from traceback import print_last
from docx import Document
from docx.table import Table
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_COLOR_INDEX
from flask import Flask, flash
from PIL import Image
import chardet


def get_image_size(image_path):
    with Image.open(image_path) as img:
        width, height = img.size
    return width, height


def process_surplus(doc, tiers, tiers_name, tiers_email,su):
    try: 

        centrale = ''
        contract_number = '10000'
        prm = ''


     

        #GET CENTRALE NAME 
        centrale = doc.paragraphs[2].text
        print(centrale)
        

        for p_index, p in enumerate(doc.paragraphs):

            if "la position du (des) Point(s) de Décompte." in p.text:

  
                p.alignment = 1
                width, height = get_image_size(su)
                if width > height :
                    p.add_run('\n')
                    p.add_run('\n')
                    p.add_run('\n')
                    p.add_run('\n')
                    p.add_run().add_picture(su, width=Inches(7), height=Inches(5))
                if width < height :
                    p.add_run('\n')
                    p.add_run('\n')
                    p.add_run('\n')
                    p.add_run().add_picture(su, width=Inches(5), height=Inches(7))


                    

           
        saved_file = "CP_"+contract_number+'_'+centrale+".docx"
        doc.save(saved_file)
        return(saved_file)
        
    except Exception as e:
            import traceback
            traceback.print_exc()
            print(f"An error occurred: {e}")