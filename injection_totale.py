# -*- coding: utf-8 -*-

from datetime import datetime
from traceback import print_last
from docx import Document
from docx.table import Table
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_COLOR_INDEX
from flask import Flask, flash
from PIL import Image
import chardet


def get_image_size(image_path):
    with Image.open(image_path) as img:
        width, height = img.size
    return width, height



def process_total(doc, tiers, tiers_name, tiers_email, su):
    
    print("this is totalité")

    try:

        centrale = ''
        contract_number = ''
        prm = ''

        

        #GET CENTRALE NAME 
        centrale = doc.paragraphs[2].text

         #GET CONTRACT NUMBER
        digits = ''
        # Iterate through each character in the string
        for char in doc.paragraphs[1].text:
        # Check if the character is a digit
            if char.isdigit():
                digits += char
        
        contract_number = digits
        doc.paragraphs[1].clear()
        doc.paragraphs[1].text = "Contrat n°"+digits+ "\n pour le site SOLAIRE PHOTOVOLTAIQUE"
        # doc.paragraphs[2].text = "pour le site SOLAIRE PHOTOVOLTAIQUE"
        run = doc.paragraphs[1].runs[0]
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(22,34,211)
        run.font.name = 'Enedis'
        run2 = doc.paragraphs[2].runs[0]
        run2.font.size = Pt(20)
        run2.font.color.rgb = RGBColor(22,34,211)
        run2.font.name = 'Enedis'
        
        # doc.paragraphs[3]._element.getparent().remove(doc.paragraphs[3]._element)
        
        #GET PRM NUMBER
        prm = doc.tables[8].rows[0].cells[1].paragraphs[0].text
        #DISTRIBUTEUR REPLACE
        doc.tables[8].rows[3].cells[1].paragraphs[0].clear()
        doc.tables[8].rows[3].cells[1].paragraphs[0].text = "Distributeur"
        doc.tables[8].rows[3].cells[1].paragraphs[0].runs[0].font.color.rgb =  RGBColor(89,89,89)
        doc.tables[8].rows[3].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        print(prm)



        for p_index, p in enumerate(doc.paragraphs):

            if "la position du (des) Point(s) de Décompte." in p.text:

  
                p.alignment = 1
                width, height = get_image_size(su)
                if width > height :
                    p.add_run('\n')
                    p.add_run('\n')
                    p.add_run('\n')
                    p.add_run('\n')
                    p.add_run().add_picture(su, width=Inches(7), height=Inches(5))
                if width < height :
                    p.add_run('\n')
                    p.add_run('\n')
                    p.add_run('\n')
                    p.add_run().add_picture(su, width=Inches(5), height=Inches(7))



        saved_file = "CP_"+contract_number+'_'+centrale+".doc"
        doc.save(saved_file)
        return(saved_file)
        
    except Exception as e:
            import traceback
            traceback.print_exc()
            print(f"An error occurred: {e}")